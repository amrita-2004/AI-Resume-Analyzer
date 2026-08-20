import re
import io
import PyPDF2

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from analyzer import SKILLS_DB, HIGH_IMPACT_VERBS, WEAK_VERBS, SECTION_HEADERS, analyze_resume_comprehensive
from jd_matcher import calculate_jd_match, extract_keywords_from_text
from ai_rewriter import rewrite_bullet_point
from interview_gen import generate_interview_questions

def extract_text_from_file(file_storage):
    """
    Extracts text from uploaded file (PDF or DOCX) in-memory.
    Returns (extracted_text, filename).
    """
    filename = file_storage.filename or "uploaded_file"
    file_bytes = file_storage.read()
    file_storage.seek(0)
    
    text = ""
    lower_name = filename.lower()
    
    if lower_name.endswith('.pdf'):
        try:
            pdf_file_like = io.BytesIO(file_bytes)
            reader = PyPDF2.PdfReader(pdf_file_like)
            for page in reader.pages:
                p_text = page.extract_text()
                if p_text:
                    text += p_text + "\n"
        except Exception as e:
            print(f"PDF extraction error: {e}")
            text = ""
    elif lower_name.endswith('.docx'):
        if HAS_DOCX:
            try:
                docx_file_like = io.BytesIO(file_bytes)
                doc = docx.Document(docx_file_like)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        full_text.append(" ".join(cell.text for cell in row.cells))
                text = "\n".join(full_text)
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                text = ""
        else:
            # Fallback simple string extraction for xml tags if docx module not available
            try:
                raw_str = file_bytes.decode('utf-8', errors='ignore')
                text = re.sub(r'<[^>]+>', ' ', raw_str)
                text = re.sub(r'\s+', ' ', text)
            except Exception:
                text = ""
    else:
        # Fallback text decoding
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
        except Exception:
            text = ""
            
    return text.strip(), filename

def analyze_job_readiness(resume_text, jd_text="", target_role=""):
    """
    Main entry point for AI Job Readiness & Rejection Recovery Engine.
    Combines comprehensive resume analysis with job description alignment.
    """
    comp_analysis = analyze_resume_comprehensive(resume_text)
    jd_analysis = calculate_jd_match(resume_text, jd_text) if jd_text and jd_text.strip() else None
    
    # 1. Calculate 7-dimensional Scores & Readiness Index
    scores = calculate_readiness_breakdown(comp_analysis, jd_analysis, resume_text, jd_text)
    readiness_score = scores["readiness_score"]
    status_tier = classify_readiness_tier(readiness_score)
    
    # 2. Rejection Vector Analysis ("Why You May Not Be Shortlisted")
    rejection_reasons = analyze_rejection_reasons(comp_analysis, jd_analysis, resume_text, jd_text)
    
    # 3. Prioritized Skill Gap Analysis
    skill_gaps = analyze_prioritized_skill_gaps(resume_text, jd_text)
    
    # 4. Personalized 7/30/60/90-Day Recovery Roadmap
    recovery_plan = generate_recovery_plan(skill_gaps, target_role or "Target Role")
    
    # 5. Targeted Project Recommendation
    project_rec = generate_project_recommendations(skill_gaps, target_role or "Target Role")
    
    # 6. Resume Improvement Suggestions (Current vs Improved)
    resume_improvements = generate_resume_improvements(resume_text, skill_gaps)
    
    # 7. Comprehensive Interview Preparation
    interview_prep = generate_interview_prep(comp_analysis, jd_analysis, skill_gaps, target_role)
    
    # 8. Re-Application Readiness Tracker Tasks
    tracker_tasks = generate_tracker_tasks(recovery_plan, project_rec, resume_improvements, skill_gaps)
    
    # 9. Final Action Recommendation
    action_rec = generate_action_recommendation(readiness_score, status_tier)
    
    return {
        "target_role": target_role or "Target Role",
        "has_jd": bool(jd_text and jd_text.strip()),
        "scores": scores,
        "readiness_score": readiness_score,
        "status_tier": status_tier,
        "disclaimer": "This readiness score is an estimated alignment metric based on parsed data and ATS criteria. It is NOT a guarantee of getting shortlisted or hired.",
        "rejection_reasons": rejection_reasons,
        "skill_gaps": skill_gaps,
        "recovery_plan": recovery_plan,
        "project_rec": project_rec,
        "resume_improvements": resume_improvements,
        "interview_prep": interview_prep,
        "tracker_tasks": tracker_tasks,
        "action_rec": action_rec,
        "comprehensive": comp_analysis,
        "jd_match": jd_analysis
    }

def calculate_readiness_breakdown(comp, jd_match, resume_text, jd_text):
    """Calculates 7 distinct scores and composite Job Readiness Score (0-100)."""
    ats_score = comp["ats_score"]
    
    if jd_match:
        overall_match = jd_match["match_percentage"]
        
        # Skill Match %
        jd_skills, _ = extract_keywords_from_text(jd_text)
        resume_skills, _ = extract_keywords_from_text(resume_text)
        
        jd_set = set([s.lower() for s in jd_skills])
        res_set = set([s.lower() for s in resume_skills])
        
        if jd_set:
            matched_count = len(jd_set.intersection(res_set))
            skills_match = int((matched_count / len(jd_set)) * 100)
        else:
            skills_match = overall_match
            
        keyword_match = jd_match["similarity_score"]
    else:
        overall_match = ats_score
        skills_match = comp["score_breakdown"]["keywords"]
        keyword_match = comp["score_breakdown"]["keywords"]
        
    # Experience Match %
    has_exp = comp["formatting_details"]["detected_sections"]["experience"]
    verb_score = comp["verb_details"]["verb_score"]
    metrics_score = comp["formatting_details"]["metrics_score"]
    experience_match = int((0.4 * (100 if has_exp else 30)) + (0.4 * verb_score) + (0.2 * metrics_score))
    experience_match = min(max(experience_match, 10), 98)
    
    # Project Relevance %
    has_projects = comp["formatting_details"]["detected_sections"]["projects"]
    project_score = 85 if has_projects else 40
    if jd_text:
        proj_match_boost = 10 if ("project" in resume_text.lower() and skills_match > 40) else 0
        project_relevance = min(project_score + proj_match_boost, 95)
    else:
        project_relevance = project_score
        
    # Education Match %
    has_edu = comp["formatting_details"]["detected_sections"]["education"]
    edu_score = 90 if has_edu else 35
    degree_found = any(deg in resume_text.lower() for deg in ["bachelor", "master", "bs", "ms", "b.tech", "m.tech", "degree", "computer science", "engineering"])
    if degree_found:
        edu_score = min(edu_score + 10, 100)
    education_match = edu_score
    
    # Weighted Composite Job Readiness Score
    if jd_match:
        readiness_score = int(
            (overall_match * 0.30) +
            (skills_match * 0.25) +
            (experience_match * 0.15) +
            (ats_score * 0.15) +
            (project_relevance * 0.10) +
            (education_match * 0.05)
        )
    else:
        readiness_score = int(
            (ats_score * 0.40) +
            (skills_match * 0.25) +
            (experience_match * 0.20) +
            (project_relevance * 0.10) +
            (education_match * 0.05)
        )
        
    readiness_score = max(min(readiness_score, 98), 15)
    
    return {
        "ats_score": ats_score,
        "overall_match": overall_match,
        "skills_match": skills_match,
        "experience_match": experience_match,
        "project_relevance": project_relevance,
        "keyword_match": keyword_match,
        "education_match": education_match,
        "readiness_score": readiness_score
    }

def classify_readiness_tier(score):
    """Classifies readiness score into 4 standard categories."""
    if score >= 85:
        return {
            "tier": "Ready to Apply",
            "range": "85–100",
            "code": "GREEN",
            "icon": "🟢",
            "badge_class": "tier-ready",
            "summary": "Your profile demonstrates strong alignment with core job requirements and high ATS readiness."
        }
    elif score >= 70:
        return {
            "tier": "Almost Ready",
            "range": "70–84",
            "code": "YELLOW_HIGH",
            "icon": "🟡",
            "badge_class": "tier-almost",
            "summary": "Your profile matches key technical skills but needs minor polish on bullet impact, metrics, or missing secondary keywords."
        }
    elif score >= 50:
        return {
            "tier": "Needs Improvement",
            "range": "50–69",
            "code": "YELLOW_LOW",
            "icon": "🟡",
            "badge_class": "tier-improvement",
            "summary": "Noticeable skill and keyword gaps detected. Practical project evidence and resume optimization recommended before applying."
        }
    else:
        return {
            "tier": "Not Ready",
            "range": "0–49",
            "code": "RED",
            "icon": "🔴",
            "badge_class": "tier-notready",
            "summary": "Significant missing requirements detected. Following the Recovery Plan and building practical projects is strongly advised."
        }

def analyze_rejection_reasons(comp, jd_match, resume_text, jd_text):
    """
    Identifies specific reasons candidate may not be shortlisted across 8 key vectors.
    Returns list of items with problem, why_it_matters, and how_to_improve.
    """
    reasons = []
    
    # Vector 1: Missing technical skills
    if jd_match and jd_match.get("missing_keywords"):
        missing_top = jd_match["missing_keywords"][:4]
        reasons.append({
            "vector": "Missing Technical Skills",
            "problem": f"Your resume lacks explicit mention of required technologies: {', '.join(missing_top)}.",
            "why_it_matters": "Recruiters and automated ATS screeners filter out candidates who do not explicitly list primary tech stack requirements.",
            "how_to_improve": f"Gain working knowledge in {missing_top[0]} through hands-on practice, then integrate it into your Skills & Projects sections."
        })
    elif len(comp["all_skills"]) < 6:
        reasons.append({
            "vector": "Insufficient Skill Density",
            "problem": f"Only {len(comp['all_skills'])} technical skills were detected on your resume.",
            "why_it_matters": "Low keyword density lowers your match confidence against target industry standards.",
            "how_to_improve": "Add a structured 'Technical Skills' section detailing languages, frameworks, databases, and DevOps tools."
        })
        
    # Vector 2: Missing job-specific keywords
    if jd_match and jd_match.get("similarity_score", 100) < 55:
        reasons.append({
            "vector": "Low Job Keyword Alignment",
            "problem": "Low vocabulary overlap between your resume and the target Job Description.",
            "why_it_matters": "ATS systems rank applicants by semantic keyword alignment. Low overlap leads to early automated rejection.",
            "how_to_improve": "Mirror exact terminology and phrases from the Job Description into your summary and project bullets."
        })
        
    # Vector 3: Weak or missing experience
    has_exp = comp["formatting_details"]["detected_sections"]["experience"]
    if not has_exp:
        reasons.append({
            "vector": "Missing Work Experience Section",
            "problem": "No clear 'Professional Experience' or 'Work History' header was detected.",
            "why_it_matters": "Hiring managers look for practical experience first to evaluate career trajectory and competence.",
            "how_to_improve": "Add a dedicated 'Work Experience' or 'Relevant Experience' section featuring internships, freelance, or academic roles."
        })
        
    # Vector 4: Weak projects
    has_projects = comp["formatting_details"]["detected_sections"]["projects"]
    if not has_projects:
        reasons.append({
            "vector": "Missing Practical Projects",
            "problem": "No dedicated 'Projects' section was detected on your resume.",
            "why_it_matters": "For tech roles, projects prove that you can build applications using the claimed tech stack.",
            "how_to_improve": "Build 2–3 full-stack or domain-specific projects and feature them under a prominent 'Projects' section."
        })
        
    # Vector 5: Missing measurable achievements
    metrics_found = comp["formatting_details"]["metrics_found"]
    if metrics_found < 2:
        reasons.append({
            "vector": "Lack of Quantifiable Metrics",
            "problem": f"Only {metrics_found} quantifiable metrics or percentage outcomes were found in your resume text.",
            "why_it_matters": "Resumes without metrics sound like job descriptions rather than proof of high performance.",
            "how_to_improve": "Use the STAR method to quantify outcomes (e.g., 'Optimized database queries, reducing API latency by 40%')."
        })
        
    # Vector 6: Poor resume structure
    struct_score = comp["formatting_details"]["structure_score"]
    if struct_score < 70:
        missing_headers = [sec.capitalize() for sec, found in comp["formatting_details"]["detected_sections"].items() if not found]
        reasons.append({
            "vector": "Non-Standard Resume Structure",
            "problem": f"Missing standard ATS section headers: {', '.join(missing_headers[:3])}.",
            "why_it_matters": "Unconventional section titles confuse ATS parsers, causing key sections to be misclassified or ignored.",
            "how_to_improve": "Use standard section headers: 'Summary', 'Technical Skills', 'Experience', 'Projects', and 'Education'."
        })
        
    # Vector 7: Weak action verbs
    weak_verbs_found = comp["verb_details"]["weak_verbs"]
    if weak_verbs_found or comp["verb_details"]["verb_score"] < 60:
        found_str = f" ('{', '.join(weak_verbs_found[:3])}')" if weak_verbs_found else ""
        reasons.append({
            "vector": "Passive Language & Weak Action Verbs",
            "problem": f"Detected passive phrasing{found_str} in your bullet points.",
            "why_it_matters": "Phrases like 'worked on' or 'responsible for' sound passive and diminish your apparent impact.",
            "how_to_improve": "Replace passive verbs with high-impact action verbs like 'Spearheaded', 'Engineered', 'Architected', or 'Scaled'."
        })
        
    # Vector 8: Insufficient evidence for required skills
    if jd_match and jd_match.get("missing_keywords"):
        reasons.append({
            "vector": "Insufficient Evidence for Role Requirements",
            "problem": "Claimed skills lack supporting context or project references in experience bullets.",
            "why_it_matters": "Interviewers verify listed skills against bullet point details to filter out inflated resumes.",
            "how_to_improve": "Explicitly mention how you applied each core skill in bullet points under your projects or experience."
        })
        
    if not reasons:
        reasons.append({
            "vector": "Minor Optimization Opportunities",
            "problem": "No critical structural flaws detected.",
            "why_it_matters": "High competition means top candidates continually refine alignment.",
            "how_to_improve": "Tailor your project bullet descriptions further to match specific metrics in candidate job descriptions."
        })
        
    return reasons

def analyze_prioritized_skill_gaps(resume_text, jd_text):
    """
    Extracts missing skills from JD (without inventing skills) and categorizes them into:
    HIGH PRIORITY (core languages/cloud/frameworks)
    MEDIUM PRIORITY (databases/secondary libraries)
    LOW PRIORITY (tools/methodologies/other)
    """
    if not jd_text or not jd_text.strip():
        res_skills, _ = extract_keywords_from_text(resume_text)
        res_lower = set([s.lower() for s in res_skills])
        
        default_high = ["Docker", "REST API"]
        default_med = ["PostgreSQL", "AWS"]
        default_low = ["Redis", "CI/CD"]
        
        return {
            "high_priority": [s for s in default_high if s.lower() not in res_lower],
            "medium_priority": [s for s in default_med if s.lower() not in res_lower],
            "low_priority": [s for s in default_low if s.lower() not in res_lower],
            "all_missing": [s for s in default_high + default_med + default_low if s.lower() not in res_lower]
        }
        
    jd_skills, _ = extract_keywords_from_text(jd_text)
    res_skills, _ = extract_keywords_from_text(resume_text)
    
    jd_set = set([s.lower() for s in jd_skills])
    res_set = set([s.lower() for s in res_skills])
    
    missing_raw = jd_set - res_set
    
    high_prio = []
    med_prio = []
    low_prio = []
    
    for skill_lower in missing_raw:
        skill_cap = skill_lower.capitalize()
        cat_found = None
        for cat, skills in SKILLS_DB.items():
            if skill_lower in [s.lower() for s in skills]:
                cat_found = cat
                break
                
        if cat_found in ["Programming Languages", "Web & Frameworks", "Cloud & DevOps"]:
            high_prio.append(skill_cap)
        elif cat_found in ["Databases & Storage", "Data Science & AI/ML"]:
            med_prio.append(skill_cap)
        else:
            low_prio.append(skill_cap)
            
    all_missing = high_prio + med_prio + low_prio
    if not high_prio and all_missing:
        high_prio = all_missing[:2]
        med_prio = all_missing[2:4]
        low_prio = all_missing[4:]
        
    return {
        "high_priority": high_prio,
        "medium_priority": med_prio,
        "low_priority": low_prio,
        "all_missing": all_missing
    }

def generate_recovery_plan(skill_gaps, target_role="Software Engineer"):
    """Generates structured 7-Day, 30-Day, 60-Day, and 90-Day action roadmaps."""
    high = skill_gaps["high_priority"]
    med = skill_gaps["medium_priority"]
    low = skill_gaps["low_priority"]
    
    skill_1 = high[0] if high else "REST API"
    skill_2 = high[1] if len(high) > 1 else (med[0] if med else "Docker")
    skill_3 = med[0] if med else "PostgreSQL"
    skill_4 = low[0] if low else "CI/CD & Cloud"
    
    return {
        "plan_7_day": {
            "title": "7-Day Sprint: Core Keyword Alignment & Immediate Fixes",
            "skill_to_learn": skill_1,
            "topic_to_study": f"Fundamentals & Syntax of {skill_1}",
            "practical_task": f"Build a mini 1-page script or endpoint demonstrating {skill_1}.",
            "recommended_project": f"Mini-{skill_1} Starter Utility",
            "resume_improvement_task": f"Add {skill_1} to your Technical Skills section and update top 2 bullet points with STAR verbs.",
            "interview_prep_task": f"Prepare 3 common technical interview questions covering {skill_1} concepts."
        },
        "plan_30_day": {
            "title": "30-Day Goal: Hands-on Skill & Project Building",
            "skill_to_learn": f"{skill_1} + {skill_2}",
            "topic_to_study": f"Integration patterns between {skill_1} and {skill_2} in modern {target_role} workflows.",
            "practical_task": f"Develop a functional web backend or microservice combining {skill_1} and {skill_2}.",
            "recommended_project": f"Containerized {skill_1} & {skill_2} Web Application",
            "resume_improvement_task": "Add a new Project section entry with 3 STAR bullet points detailing architecture and metric impact.",
            "interview_prep_task": f"Practice explaining the design trade-offs of {skill_2} in a 5-minute technical demo presentation."
        },
        "plan_60_day": {
            "title": "60-Day Goal: Advanced Architecture & Database Mastery",
            "skill_to_learn": f"{skill_3} & Production Optimization",
            "topic_to_study": f"Database indexing, caching strategies, and API security for {target_role} applications.",
            "practical_task": f"Integrate {skill_3} with authentication, unit testing, and logging.",
            "recommended_project": f"Scalable Data Management System featuring {skill_3}",
            "resume_improvement_task": "Quantify outcomes across all projects (e.g. latency reductions, test coverage percentages).",
            "interview_prep_task": "Conduct 2 mock technical interviews focusing on system design and database query optimization."
        },
        "plan_90_day": {
            "title": "90-Day Mastery: Cloud Deployment & Interview Readiness",
            "skill_to_learn": f"{skill_4} & Production Cloud Infrastructure",
            "topic_to_study": "CI/CD pipelines, Docker containerization, AWS/GCP deployment, and cloud monitoring.",
            "practical_task": "Deploy your complete full-stack project live with automated CI/CD and domain routing.",
            "recommended_project": f"Production-grade Live Deployed Application with {skill_1}, {skill_2}, and {skill_4}",
            "resume_improvement_task": "Perform a final Resume Re-Analysis to verify job match improvement before target applications.",
            "interview_prep_task": "Complete behavioral STAR stories and mock interview rounds with peer reviews."
        }
    }

def generate_project_recommendations(skill_gaps, target_role="Software Engineer"):
    """
    Recommends a practical, realistic project based on missing high/medium priority skills.
    Explains why useful, missing skills demonstrated, and target resume section.
    """
    missing_skills = skill_gaps["all_missing"]
    
    if missing_skills:
        primary_skills = ", ".join(missing_skills[:3])
        skills_demo = missing_skills[:3]
    else:
        primary_skills = "Flask, REST API, Docker, PostgreSQL"
        skills_demo = ["Flask", "REST API", "Docker", "PostgreSQL"]
        
    title = f"Containerized {target_role} Service with {primary_skills}"
    description = f"Build and deploy a robust backend/full-stack service utilizing {primary_skills}. Implement REST APIs, database persistence, authentication, and Docker container orchestration."
    why_useful = f"Directly addresses your target Job Description gaps by providing verifiable proof of hands-on competence in {primary_skills}."
    resume_section = "Key Projects Section (Title: " + title + ")"
    
    return {
        "title": title,
        "description": description,
        "why_useful": why_useful,
        "missing_skills_demonstrated": skills_demo,
        "resume_section_to_improve": resume_section,
        "architecture_stack": primary_skills
    }

def generate_resume_improvements(resume_text, skill_gaps):
    """
    Provides Current vs Improved bullet point recommendations without fabricating fake experience.
    """
    lines = [line.strip(' •-*') for line in resume_text.split('\n') if len(line.strip()) > 15]
    sample_bullets = lines[:3] if lines else ["Worked on web application development and backend services."]
    
    improvements = []
    for bullet in sample_bullets:
        rw = rewrite_bullet_point(bullet)
        improvements.append({
            "current": rw["original"],
            "improved": rw["rewritten"],
            "enhancement_note": "Added high-impact action verb, structure, and STAR metric framework."
        })
        
    return improvements

def generate_interview_prep(comp, jd_match, skill_gaps, target_role="Software Engineer"):
    """Generates comprehensive technical, HR, project, and skill-specific questions."""
    extracted = comp["all_skills"] + skill_gaps["all_missing"]
    questions = generate_interview_questions(extracted, target_role or "Software Engineer")
    
    pre_apply_checklist = [
        "Ensure all listed tech skills are backed by concrete project bullets.",
        "Quantify at least 3 bullet points with measurable outcomes (%, $, time saved).",
        "Verify contact details (Email, Phone, LinkedIn, GitHub) are visible in top header.",
        "Re-run Resume Re-Analysis to ensure readiness score is >= 70%."
    ]
    
    return {
        "technical_questions": questions.get("technical_questions", []),
        "behavioral_questions": questions.get("behavioral_questions", []),
        "resume_probing_questions": questions.get("resume_probing_questions", []),
        "pre_apply_checklist": pre_apply_checklist
    }

def generate_tracker_tasks(plan, project_rec, improvements, skill_gaps):
    """Generates weighted interactive progress tracker tasks across 5 stages."""
    tasks = [
        {
            "id": "task_1",
            "phase": "Skill Acquisition",
            "title": f"Study & Learn Core Gap: {plan['plan_7_day']['skill_to_learn']}",
            "weight": 15,
            "completed": False
        },
        {
            "id": "task_2",
            "phase": "Project Building",
            "title": f"Build Recommended Project: {project_rec['title']}",
            "weight": 25,
            "completed": False
        },
        {
            "id": "task_3",
            "phase": "Resume Optimization",
            "title": "Incorporate STAR Action Verbs and Metrics into Bullet Points",
            "weight": 20,
            "completed": False
        },
        {
            "id": "task_4",
            "phase": "Interview Prep",
            "title": "Prepare Technical Answers & STAR Behavioral Stories",
            "weight": 20,
            "completed": False
        },
        {
            "id": "task_5",
            "phase": "Re-Application Verification",
            "title": "Re-Analyze Resume to confirm score improvement",
            "weight": 20,
            "completed": False
        }
    ]
    return tasks

def generate_action_recommendation(readiness_score, tier_info):
    """Generates the final action recommendation banner data."""
    code = tier_info["code"]
    
    if code == "RED":
        status_text = "🔴 NOT READY"
        message = "Improve these critical areas before applying. Focus on building the recommended project and adding missing core technical keywords."
        action_type = "danger"
    elif code in ["YELLOW_LOW", "YELLOW_HIGH"]:
        status_text = "🟡 ALMOST READY"
        message = "You are close! Complete your remaining action items and polish your resume bullets before submitting your application."
        action_type = "warning"
    else:
        status_text = "🟢 READY TO APPLY"
        message = "Your profile currently demonstrates strong alignment with major requirements. You can start applying while continuing fine-tuning."
        action_type = "success"
        
    return {
        "status_text": status_text,
        "message": message,
        "action_type": action_type,
        "tier_name": tier_info["tier"],
        "safe_language_notice": "Disclaimer: 'Readiness' signifies estimated keyword alignment & structural optimization. It does NOT guarantee job offers."
    }

def compare_resumes(old_readiness, new_resume_text, jd_text="", target_role=""):
    """
    Compares a new/improved resume against a previous analysis.
    Returns delta metrics, skills gained, added keywords, and new readiness score.
    """
    new_analysis = analyze_job_readiness(new_resume_text, jd_text, target_role)
    
    old_score = old_readiness.get("readiness_score", 50)
    new_score = new_analysis["readiness_score"]
    improvement_pct = new_score - old_score
    
    old_skills = set([s.lower() for s in old_readiness.get("comprehensive", {}).get("all_skills", [])])
    new_skills = set([s.lower() for s in new_analysis["comprehensive"]["all_skills"]])
    
    gained_skills = [s.capitalize() for s in (new_skills - old_skills)]
    
    old_missing = set([s.lower() for s in old_readiness.get("skill_gaps", {}).get("all_missing", [])])
    new_missing = set([s.lower() for s in new_analysis["skill_gaps"]["all_missing"]])
    
    resolved_gaps = [s.capitalize() for s in (old_missing - new_missing)]
    
    return {
        "previous_score": old_score,
        "current_score": new_score,
        "improvement_pct": improvement_pct,
        "improvement_sign": "+" if improvement_pct >= 0 else "",
        "gained_skills": gained_skills,
        "resolved_gaps": resolved_gaps,
        "remaining_gaps": [s.capitalize() for s in new_missing],
        "new_readiness_status": new_analysis["status_tier"]["tier"],
        "new_analysis": new_analysis
    }
